import unittest
import os
import tempfile
from unittest.mock import Mock, patch

from docx import Document

from rfp.agents.generation.implementation import (
    _fit_generation_prompt,
    _merge_section_drafts,
    _normalize_batch_headings,
    _proposal_structure,
    _section_batches,
    _split_batch_sections,
    generation_agent,
)
from rfp.orchestration.routing import after_generation
from rfp.agents.extraction.implementation import (
    _extract_template_sections,
    _merge_template_outline,
)
from rfp.agents.quality.implementation import (
    _check_section_order,
    _check_template_compliance,
    _evaluate_grounding_and_coherence,
    _extract_review_json,
    _identify_failed_sections,
    _review_groups,
    _template_sections,
    quality_agent,
)


class ResponseTemplateQualityTests(unittest.TestCase):
    def test_bilingual_short_heading_is_restored_to_exact_template_title(self):
        expected = "5. Plan de travail et calendrier / Work Plan and Timeline"
        draft, unmatched = _normalize_batch_headings(
            "## 5. Plan de travail et calendrier\nContenu détaillé.",
            [expected],
        )

        self.assertIn(f"## {expected}", draft)
        self.assertEqual(unmatched, [])

    def test_numbered_reworded_heading_is_restored_to_exact_template_title(self):
        expected = "10. Proposition financière / Financial Proposal"
        draft, unmatched = _normalize_batch_headings(
            "## 10. Budget et tarification\nMontant fondé sur le cahier des charges.",
            [expected],
        )

        self.assertIn(f"## {expected}", draft)
        self.assertEqual(unmatched, [])

    def test_generated_batch_is_split_into_live_template_sections(self):
        sections = ["1. Introduction", "2. Méthodologie / Methodology"]
        content = _split_batch_sections(
            "## 1. Introduction\nContexte vérifié.\n\n"
            "## 2. Méthodologie / Methodology\nApproche proposée.",
            sections,
        )

        self.assertIn("Contexte vérifié.", content[sections[0]])
        self.assertIn("Approche proposée.", content[sections[1]])

    def test_quality_maps_a_claim_to_only_its_template_section(self):
        sections = ["1. Context", "2. Solution", "3. Schedule"]
        claim = "The proposed platform guarantees unlimited availability."
        draft = (
            "## 1. Context\nTender facts.\n\n"
            f"## 2. Solution\n{claim}\n\n"
            "## 3. Schedule\nThirty-two week schedule."
        )

        failed = _identify_failed_sections(
            draft=draft,
            sections=sections,
            missing_sections=[],
            out_of_order_sections=[],
            quality_findings={},
            grounding_review={
                "groundedness_score": 0.8,
                "coherence_score": 0.9,
                "unsupported_claims": [],
                "contradictions": [{"claim": claim, "evidence": "unsupported"}],
            },
            word_count=200,
        )

        self.assertEqual(failed, ["2. Solution"])

    def test_retry_regenerates_only_failed_section_and_preserves_others(self):
        sections = ["1. Context", "2. Solution", "3. Schedule"]
        previous_draft = (
            "## 1. Context\nAccepted context.\n\n"
            "## 2. Solution\nUnsupported guarantee.\n\n"
            "## 3. Schedule\nAccepted schedule."
        )
        prior_batches = [
            {
                "sections": [section],
                "draft": block,
                "tender_excerpts": "evidence",
                "response_template_excerpts": "instructions",
            }
            for section, block in _split_batch_sections(
                previous_draft, sections
            ).items()
        ]
        rag = Mock()
        rag.query.return_value = "relevant evidence"
        knowledge = Mock()
        provider = Mock()
        provider.complete.return_value = (
            "## 2. Solution\nRepaired solution grounded in tender evidence."
        )

        with patch(
            "rfp.agents.generation.implementation.get_provider",
            return_value=provider,
        ):
            result = generation_agent(
                {
                    "is_verified": True,
                    "workspace_slug": "tender",
                    "response_template_workspace_slug": "template",
                    "requirements": {
                        "scope_summary": "digital platform",
                        "response_template": {
                            "required_sections": sections,
                            "section_order": sections,
                        },
                    },
                    "research_summary": "relevant research",
                    "previous_draft": previous_draft,
                    "previous_generation_evidence": {
                        "section_batches": prior_batches,
                        "project_references": "references",
                        "cv_excerpts": "profiles",
                        "past_proposals": "proposals",
                    },
                    "generation_attempts": 1,
                    "quality_report": {"failed_sections": ["2. Solution"]},
                },
                rag=rag,
                knowledge=knowledge,
            )

        self.assertEqual(provider.complete.call_count, 1)
        self.assertIn("Accepted context.", result["draft_proposal"])
        self.assertIn("Repaired solution grounded", result["draft_proposal"])
        self.assertIn("Accepted schedule.", result["draft_proposal"])
        self.assertNotIn("Unsupported guarantee.", result["draft_proposal"])
        self.assertTrue(result["generation_evidence"]["repair_mode"])
        self.assertEqual(
            result["generation_evidence"]["repaired_sections"],
            ["2. Solution"],
        )

    def test_quality_evaluator_falls_back_when_groq_rejects_json_mode(self):
        provider = Mock()
        provider.complete.side_effect = [
            RuntimeError(
                "Groq completion failed: HTTP 400: "
                '{"error":{"code":"json_validate_failed"}}'
            ),
            """{
              "groundedness_score": 0.84,
              "coherence_score": 0.88,
              "unsupported_claims": [],
              "contradictions": [],
              "coherence_issues": [],
              "notes": []
            }""",
        ]

        with patch("rfp.agents.quality.implementation.get_provider", return_value=provider):
            review = _evaluate_grounding_and_coherence(
                {"generation_evidence": {"requirements": {"scope": "Test"}}},
                "# Proposal\nGrounded draft content.",
            )

        self.assertEqual(provider.complete.call_count, 2)
        self.assertEqual(review["groundedness_score"], 0.84)
        self.assertEqual(review["coherence_score"], 0.88)
        self.assertNotIn("evaluation_error", review)

    def test_quality_review_repairs_truncated_json_with_trailing_comma(self):
        malformed = """```json
        {
          "groundedness_score": 0.82,
          "coherence_score": 0.91,
          "unsupported_claims": [],
          "contradictions": [],
          "coherence_issues": [],
          "notes": ["Evaluation completed",]
        """

        review = _extract_review_json(malformed)

        self.assertEqual(review["groundedness_score"], 0.82)
        self.assertEqual(review["coherence_score"], 0.91)
        self.assertEqual(review["notes"], ["Evaluation completed"])

    def test_docx_heading_structure_recovers_complete_template_outline(self):
        document = Document()
        expected = [
            "1. Introduction",
            "2. Compréhension du besoin",
            "3. Conformité aux exigences",
            "4. Approche et méthodologie",
            "5. Équipe et organisation",
            "6. Calendrier du projet",
            "7. Budget et tarification",
            "8. Assurance qualité",
            "9. Gestion des risques",
            "10. Acceptation et garantie",
            "11. Annexes",
        ]
        for heading in expected:
            document.add_heading(heading, level=1)
            document.add_paragraph("Instructions for this section.")
        handle, path = tempfile.mkstemp(suffix=".docx")
        os.close(handle)
        try:
            document.save(path)
            self.assertEqual(_extract_template_sections(path), expected)
        finally:
            os.unlink(path)

    def test_complete_local_outline_replaces_partial_rag_outline(self):
        requirements = {
            "response_template": {
                "required_sections": ["2. Need", "3. Compliance"],
                "section_order": ["2. Need", "3. Compliance"],
                "instructions": ["Keep this instruction"],
            }
        }
        complete = ["1. Introduction", "2. Need", "3. Compliance", "4. Annexes"]

        merged = _merge_template_outline(requirements, complete)

        self.assertEqual(merged["response_template"]["section_order"], complete)
        self.assertEqual(
            merged["response_template"]["instructions"],
            ["Keep this instruction"],
        )
        self.assertEqual(
            merged["response_template"]["outline_source"],
            "local_document_structure",
        )

    def test_generation_prompt_enforces_total_budget_without_cutting_instructions(self):
        huge = "French tender evidence and requirements. " * 1000
        prompt, fitted = _fit_generation_prompt(
            {
                "batch_number": 1,
                "batch_count": 2,
                "tender_excerpts": huge,
                "response_template_excerpts": huge,
                "response_template_rules": huge,
                "proposal_structure": "## 1. Introduction\n## 2. Methodology",
                "revision_feedback": huge,
                "requirements": huge,
                "research_summary": huge,
                "project_references": huge,
                "cv_excerpts": huge,
                "past_proposals": huge,
            },
            max_chars=13000,
        )

        self.assertLessEqual(len(prompt), 13000)
        self.assertIn("## 1. Introduction", prompt)
        self.assertIn("Do not invent specific figures", prompt)
        self.assertLess(len(fitted["research_summary"]), len(huge))

    def test_quality_failure_does_not_regenerate_by_default(self):
        sections = ["Contexte", "Solution", "Planning"]
        draft = "\n".join(
            f"# {section}\n" + ("substantive content " * 60)
            for section in sections
        )
        evaluator_result = {
            "groundedness_score": 0.5,
            "coherence_score": 0.9,
            "unsupported_claims": [{"claim": "x", "reason": "unsupported"}],
            "contradictions": [],
            "coherence_issues": [],
            "notes": [],
        }

        with patch(
            "rfp.agents.quality.implementation._evaluate_grounding_and_coherence",
            return_value=evaluator_result,
        ):
            result = quality_agent(
                {
                    "is_verified": True,
                    "security_passed": True,
                    "draft_proposal": draft,
                    "generation_attempts": 1,
                    "requirements": {
                        "response_template": {
                            "required_sections": sections,
                            "section_order": sections,
                        }
                    },
                }
            )

        self.assertEqual(result["status"], "failed")

    def test_quality_review_uses_one_compact_group_by_default(self):
        section_batches = [
            {"sections": [f"S{index}"], "draft": f"draft {index}", "evidence": index}
            for index in range(4)
        ]
        groups = _review_groups(
            {"generation_evidence": {"section_batches": section_batches}},
            "full draft",
        )

        self.assertEqual(len(groups), 1)
        self.assertIn("draft 0", groups[0][1])
        self.assertIn("draft 1", groups[0][1])
        self.assertIn("draft 2", groups[0][1])
        self.assertNotIn("draft", groups[0][0]["section_batches"][0])

    def test_quality_evaluator_error_does_not_retry_generation(self):
        draft = "\n".join(
            f"# {section}\n" + ("substantive content " * 30)
            for section in [
                "Executive Summary",
                "Understanding of the Requirements",
                "Proposed Approach & Methodology",
                "Indicative Work Plan / Timeline",
                "Risk Management & Quality Assurance",
                "Proposed Team (Profils Proposés)",
                "Why Us",
            ]
        )
        evaluator_result = {
            "groundedness_score": 0.0,
            "coherence_score": 0.0,
            "unsupported_claims": [],
            "contradictions": [],
            "coherence_issues": [],
            "notes": [],
            "evaluation_error": "HTTP 413 request too large",
        }

        with patch(
            "rfp.agents.quality.implementation._evaluate_grounding_and_coherence",
            return_value=evaluator_result,
        ):
            result = quality_agent(
                {
                    "is_verified": True,
                    "security_passed": True,
                    "draft_proposal": draft,
                    "generation_attempts": 1,
                }
            )

        self.assertEqual(result["status"], "failed")
        self.assertFalse(result["quality_passed"])
        self.assertFalse(result["quality_report"]["evaluation_available"])
        self.assertFalse(
            any(
                "groundedness=0.00" in note
                for note in result["quality_report"]["notes"]
            )
        )
        self.assertTrue(
            any("without regenerating" in note for note in result["quality_report"]["notes"])
        )

    def test_template_sections_are_batched_dynamically_in_groups_of_three(self):
        sections = [f"Custom section {index}" for index in range(1, 8)]

        batches = _section_batches(
            {"section_order": sections},
            batch_size=3,
        )

        self.assertEqual(
            batches,
            [sections[0:3], sections[3:6], sections[6:7]],
        )

    def test_empty_generation_stops_before_security(self):
        self.assertNotEqual(after_generation({"generation": {"draft_proposal": ""}}), "security")
        self.assertEqual(
            after_generation({"generation": {"draft_proposal": "proposal"}}), "security"
        )

    def test_client_template_replaces_default_generation_outline(self):
        outline = _proposal_structure(
            {
                "required_sections": ["1. Introduction", "2. Compréhension du besoin"],
                "section_order": ["1. Introduction", "2. Compréhension du besoin"],
            }
        )

        self.assertIn("## 1. Introduction", outline)
        self.assertIn("## 2. Compréhension du besoin", outline)
        self.assertNotIn("Executive Summary", outline)

    def test_client_template_sections_override_defaults(self):
        state = {
            "requirements": {
                "response_template": {
                    "required_sections": ["Contexte", "Solution", "Planning"],
                    "section_order": ["Contexte", "Solution", "Planning"],
                }
            }
        }

        required, ordered = _template_sections(state)

        self.assertEqual(required, ["Contexte", "Solution", "Planning"])
        self.assertEqual(ordered, required)

    def test_missing_and_out_of_order_sections_are_reported(self):
        draft = "# Solution\nDetails\n# Contexte\nDetails"
        sections = ["Contexte", "Solution", "Planning"]

        self.assertEqual(_check_template_compliance(draft, sections), ["Planning"])
        self.assertEqual(_check_section_order(draft, sections), sections)

    def test_numbered_template_title_matches_unnumbered_markdown_heading(self):
        draft = "# Introduction\nTexte\n## **2. Compréhension du besoin**\nTexte"
        sections = ["1. Introduction", "2. Compréhension du besoin"]

        self.assertEqual(_check_template_compliance(draft, sections), [])
        self.assertEqual(_check_section_order(draft, sections), [])


if __name__ == "__main__":
    unittest.main()
